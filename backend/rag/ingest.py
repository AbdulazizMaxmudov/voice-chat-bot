"""
documents/ papkasidagi Word fayllarni o'qib, ChromaDB ga yuklash moduli.
Gemini embedding-001 modeli bilan vektorlashtiradi.
"""

import os
import logging
from pathlib import Path

import chromadb
import google.generativeai as genai
from docx import Document

# Yo'llar: bu fayl backend/rag/ da, shuning uchun ikki daraja yuqorida backend/
BASE_DIR = Path(__file__).resolve().parent.parent
DOCUMENTS_DIR = BASE_DIR / "documents"
CHROMA_DIR = BASE_DIR / "chroma_db"
COLLECTION_NAME = "documents"

logger = logging.getLogger(__name__)


def _setup_gemini() -> None:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY .env faylida topilmadi")
    genai.configure(api_key=api_key)


def _read_docx(file_path: Path) -> str:
    """Word fayldan paragraflar va jadvallarni o'qiydi."""
    doc = Document(file_path)
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _chunk_text(text: str, chunk_size: int = 1500, overlap: int = 300) -> list[str]:
    """
    Paragraph-aware chunking: paragraflar bo'yicha kesadi, jumlaning o'rtasida to'xtamaydi.
    Yagona paragraf chunk_size dan katta bo'lsa — belgi bo'yicha kesiladi.
    """
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para)

        # Joriy chunk to'ldi — saqlash
        if current_len + para_len + 1 > chunk_size and current:
            chunks.append("\n".join(current))
            # Overlap: oxirgi bir necha paragrafni keyingi chunkga olib o'tish
            overlap_buf: list[str] = []
            overlap_chars = 0
            for p in reversed(current):
                if overlap_chars + len(p) <= overlap:
                    overlap_buf.insert(0, p)
                    overlap_chars += len(p)
                else:
                    break
            current = overlap_buf
            current_len = sum(len(p) for p in current)

        # Bitta paragraf juda katta bo'lsa — belgi bo'yicha kesish
        if para_len > chunk_size:
            if current:
                chunks.append("\n".join(current))
                current = []
                current_len = 0
            start = 0
            while start < para_len:
                end = min(start + chunk_size, para_len)
                chunks.append(para[start:end])
                next_start = end - overlap
                start = next_start if next_start > start else end
        else:
            current.append(para)
            current_len += para_len + 1

    if current:
        chunks.append("\n".join(current))

    return [c for c in chunks if c.strip()]


def _embed_batch_with_retry(batch: list[str], batch_num: int, total: int) -> list[list[float]]:
    """Bitta batchni embed qiladi; 429 bo'lsa qayta urinadi."""
    import time

    wait_times = [15, 30, 65]
    for attempt, wait in enumerate(wait_times + [None], start=1):
        try:
            result = genai.embed_content(
                model="models/gemini-embedding-001",
                content=batch,
                task_type="retrieval_document",
            )
            logger.info(f"Batch {batch_num}/{total} tayyor: {len(batch)} ta chunk")
            return result["embedding"]
        except Exception as e:
            if "429" in str(e) and wait is not None:
                logger.warning(f"Rate limit (429), {wait}s kutilmoqda (urinish {attempt}/{len(wait_times)})...")
                time.sleep(wait)
            else:
                raise
    raise RuntimeError("Embedding uchun barcha urinishlar tugadi")


def _embed_texts(texts: list[str]) -> list[list[float]]:
    """
    Matnlarni 100 tadan batch qilib embed qiladi.
    429 xato faqat _embed_batch_with_retry ichida retry orqali boshqariladi.
    """
    BATCH_SIZE = 100
    all_embeddings: list[list[float]] = []
    total_batches = (len(texts) + BATCH_SIZE - 1) // BATCH_SIZE

    for i in range(0, len(texts), BATCH_SIZE):
        batch = texts[i : i + BATCH_SIZE]
        batch_num = i // BATCH_SIZE + 1
        embeddings = _embed_batch_with_retry(batch, batch_num, total_batches)
        all_embeddings.extend(embeddings)

    return all_embeddings


def _ingest_sync() -> dict:
    """Blokirovchi ingest — run_in_executor ichida ishlaydi."""
    _setup_gemini()

    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

    docx_files = [f for f in DOCUMENTS_DIR.glob("*.docx") if not f.name.startswith("~$")]
    if not docx_files:
        logger.info("documents/ papkasida .docx fayl topilmadi")
        return {"message": "Hech qanday fayl topilmadi"}

    client = chromadb.PersistentClient(path=str(CHROMA_DIR))

    try:
        client.delete_collection(COLLECTION_NAME)
        logger.info("Eski collection o'chirildi")
    except Exception:
        pass

    collection = client.create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    chunk_offset = 0
    loaded_files: list[str] = []

    for file_path in docx_files:
        try:
            logger.info(f"O'qilmoqda: {file_path.name}")
            text = _read_docx(file_path)
            if not text.strip():
                logger.warning(f"{file_path.name} bo'sh, o'tkazildi")
                continue

            chunks = _chunk_text(text)
            if not chunks:
                continue

            logger.info(f"{file_path.name}: {len(chunks)} ta chunk")
            embeddings = _embed_texts(chunks)

            collection.add(
                documents=chunks,
                embeddings=embeddings,
                ids=[f"chunk_{chunk_offset + i}" for i in range(len(chunks))],
                metadatas=[{"source": file_path.name} for _ in chunks],
            )

            chunk_offset += len(chunks)
            loaded_files.append(file_path.name)
            logger.info(f"{file_path.name} muvaffaqiyatli yuklandi")

        except Exception as e:
            logger.error(f"{file_path.name} yuklashda xato: {e}")

    if not loaded_files:
        return {"message": "Fayllar o'qishda xatolik yuz berdi, loglarni tekshiring"}

    return {
        "message": f"{len(loaded_files)} ta fayl yuklandi",
        "files": loaded_files,
    }


async def ingest_documents() -> dict:
    """
    documents/ papkasidagi barcha .docx fayllarni ChromaDB ga yuklaydi.
    Blokirovchi ish thread pool da bajariladi — server muzlamaydi.
    """
    import asyncio
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, _ingest_sync)
